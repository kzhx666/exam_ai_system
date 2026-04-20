import streamlit as st
import os
import requests
import docx
from docx.shared import Pt, RGBColor, Mm, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import re
import uuid
import fitz
import subprocess
import hashlib
from PIL import Image, ImageOps, ImageFilter
import numpy as np
import google.generativeai as genai
from dotenv import load_dotenv
import sqlite3
from datetime import datetime
import importlib

# ================= 依赖热加载 =================
def ensure_dependencies():
    missing = []
    for pkg in ['pandas', 'openpyxl']:
        try:
            importlib.import_module(pkg)
        except ImportError:
            missing.append(pkg)
    if missing:
        st.info(f"正在自动安装必要组件 ({', '.join(missing)})... 请稍候。")
        subprocess.check_call(["pip", "install", *missing])
        st.success("组件安装完成！")

ensure_dependencies()
import pandas as pd
# ==============================================

load_dotenv()

# ================= UI 全局高级美化配置 =================
st.set_page_config(page_title="AI 试卷排版大师 v2.9", layout="wide", page_icon="📝")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container { padding-top: 2rem; padding-bottom: 2rem; }
    .stButton>button { border-radius: 8px !important; font-weight: 600 !important; transition: all 0.3s ease !important; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 4px 12px rgba(0,0,0,0.1); }
    div[data-testid="stExpander"] { background-color: #ffffff; border-radius: 10px !important; box-shadow: 0 2px 8px rgba(0,0,0,0.04) !important; border: 1px solid #f0f2f6 !important; margin-bottom: 1rem; }
    div[data-testid="stExpander"] > summary { font-weight: bold; font-size: 1.1rem; }
    [data-testid="stSidebar"] { background-color: #f8f9fa; border-right: 1px solid #e9ecef; }
    
    [data-testid="stFileUploadDropzone"] > div > div > span { display: none; }
    [data-testid="stFileUploadDropzone"] > div > div::before { content: "将试卷文件拖拽至此处"; display: block; font-size: 16px; font-weight: 500; margin-bottom: 5px; }
    [data-testid="stFileUploadDropzone"] > div > div > small { display: none; }
    [data-testid="stFileUploadDropzone"] > div > div::after { content: "文件大小上限 200MB • 支持 PDF, DOCX, DOC, PNG, JPG, TXT"; display: block; font-size: 13px; opacity: 0.6; margin-top: 5px; }
    [data-testid="stFileUploadDropzone"] button { color: transparent !important; }
    [data-testid="stFileUploadDropzone"] button::after { content: "浏览本地文件"; color: currentColor; position: absolute; left: 50%; top: 50%; transform: translate(-50%, -50%); font-weight: 400; }
</style>
""", unsafe_allow_html=True)
# =======================================================

DATA_DIR = "/app/data"
IMAGE_DIR = os.path.join(DATA_DIR, "images")
DB_PATH = os.path.join(DATA_DIR, "exam_bank.db")
os.makedirs(IMAGE_DIR, exist_ok=True)

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS questions
                 (id TEXT PRIMARY KEY, type TEXT, section TEXT, content TEXT, options TEXT, answer TEXT, explanation TEXT, tags TEXT, add_time TEXT)''')
    conn.commit()
    return conn

conn = init_db()

def save_to_db(questions, tags=""):
    c = conn.cursor()
    count = 0
    for q in questions:
        q_id = str(uuid.uuid4())[:8]
        opts_str = json.dumps(q.get('options', []), ensure_ascii=False)
        try:
            c.execute("INSERT INTO questions VALUES (?,?,?,?,?,?,?,?,?)",
                      (q_id, q.get('type',''), q.get('section',''), q.get('content',''), opts_str, q.get('answer',''), q.get('explanation',''), tags, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            count += 1
        except: pass
    conn.commit()
    return count

def load_from_db(tag_filter=None):
    c = conn.cursor()
    if tag_filter:
        c.execute("SELECT * FROM questions WHERE tags LIKE ?", ('%'+tag_filter+'%',))
    else:
        c.execute("SELECT * FROM questions ORDER BY add_time DESC LIMIT 100")
    rows = c.fetchall()
    res = []
    for r in rows:
        res.append({"id": r[0], "type": r[1], "section": r[2], "content": r[3], "options": json.loads(r[4]), "answer": r[5], "explanation": r[6], "tags": r[7]})
    return res

if 'cart' not in st.session_state:
    st.session_state.cart = []
if 'section_order' not in st.session_state:
    st.session_state.section_order = []

with st.sidebar:
    st.image("https://img.shields.io/badge/AI%20排版引擎-v2.9%20Pro-blue?style=for-the-badge", use_column_width=True)
    st.header("🗄️ 树状题库与收藏")
    tag_search = st.text_input("🔍 搜标签 (如: 建筑材料/第一章)", placeholder="输入知识点标签搜索库中试题")
    if st.button("从本地库拉取至手术台"):
        db_qs = load_from_db(tag_search)
        if db_qs:
            st.session_state.cart.extend(db_qs)
            st.success(f"已拉取 {len(db_qs)} 题！")
        else:
            st.warning("未找到匹配题目。")
            
    st.divider()
    st.header("🖼️ 漏图补丁站")
    patch_files = st.file_uploader("上传缺失的图片", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="patch_uploader", label_visibility="collapsed")
    if patch_files:
        for pf in patch_files:
            file_bytes = pf.read()
            file_hash = hashlib.md5(file_bytes).hexdigest()[:6]
            img_ext = pf.name.split('.')[-1].lower()
            safe_name = f"patch_{file_hash}.{img_ext}"
            img_path = os.path.join(IMAGE_DIR, safe_name)
            if not os.path.exists(img_path):
                with open(img_path, "wb") as f:
                    f.write(file_bytes)
            st.image(img_path, width=150)
            st.code(f"[图片:{safe_name}]", language="text")

st.title("📝 建筑材料教务系统 - 全能排版矩阵")
st.caption("✨ v2.9：国标级括弧排版间距 | 智能题型归一化 | 终极洗白防丢")

def normalize_type(raw_type):
    t = str(raw_type).lower()
    if 'single' in t or '单选' in t: return '单选题'
    if 'multiple' in t or '多选' in t: return '多选题'
    if 'judge' in t or '判断' in t or 'true' in t or 'false' in t or 't/f' in t: return '判断题'
    if 'fill' in t or '填空' in t or 'blank' in t: return '填空题'
    if 'short' in t or '简答' in t or 'answer' in t: return '简答题'
    if 'calc' in t or '计算' in t: return '计算题'
    if 'draw' in t or '绘图' in t or '画图' in t: return '绘图题'
    return raw_type

def clean_option(opt_str):
    return re.sub(r'^[\(（\[]?\s*[A-HＡ-Ｈ]\s*([\)）\]\.\、\．\:]+|\s+)\s*', '', str(opt_str), flags=re.IGNORECASE).strip()

def clean_stem(text):
    return str(text).strip()

def calculate_option_layout(question):
    q_type = normalize_type(question.get('type', ''))
    if q_type in ['判断题', '填空题', '简答题', '计算题', '绘图题']: return "none"
    options_list = question.get('options', [])
    if not options_list or len(options_list) == 0: return "none"
    if len(options_list) != 4: return "block"
    clean_opts = [clean_option(opt) for opt in options_list]
    max_length = max(len(opt) for opt in clean_opts)
    if max_length <= 5: return "inline"
    elif max_length <= 15: return "grid"
    else: return "block"

def group_questions(cart_list, section_configs, ordered_sections):
    groups = {}
    for q in cart_list:
        sec = q.get('clean_section', '未命名板块')
        if sec not in groups: groups[sec] = []
        groups[sec].append(q)
    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    sections = []
    for i, sec_name in enumerate(ordered_sections):
        if sec_name not in groups: continue
        qs = groups[sec_name]
        pts = section_configs.get(sec_name, {}).get('score', 0)
        sec_score = len(qs) * pts
        if re.match(r'^([一二三四五六七八九十]+|[0-9]+)[\、\.\s]+', sec_name):
            title_text = f"{sec_name}（本大题共 {len(qs)} 小题，每小题 {pts:g} 分，共 {sec_score:g} 分）"
        else:
            idx = cn_nums[i] if i < len(cn_nums) else str(i+1)
            title_text = f"{idx}、{sec_name}（本大题共 {len(qs)} 小题，每小题 {pts:g} 分，共 {sec_score:g} 分）"
        sections.append({"title": title_text, "questions": qs})
    return sections

def extract_text_and_images(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    file_bytes = uploaded_file.getvalue()
    if ext == 'txt': return file_bytes.decode('utf-8', errors='ignore')
    elif ext == 'pdf':
        text_content = []
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                blocks = page.get_text("dict")["blocks"]
                blocks.sort(key=lambda b: b["bbox"][1]) 
                for b in blocks:
                    if b["type"] == 0:  
                        para_text = ""
                        for line in b["lines"]:
                            for span in line["spans"]: para_text += span["text"]
                        if para_text.strip(): text_content.append(para_text)
                    elif b["type"] == 1: 
                        img_ext = b.get("ext", "png")
                        if b.get("image"):
                            img_name = f"img_{uuid.uuid4().hex[:6]}.{img_ext}"
                            with open(os.path.join(IMAGE_DIR, img_name), "wb") as f: f.write(b.get("image"))
                            text_content.append(f"\n[图片:{img_name}]\n")
            return "\n".join(text_content)
        except: return None
    elif ext == 'docx':
        text_content = []
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    para = docx.text.paragraph.Paragraph(element, doc)
                    para_text = "".join(run.text for run in para.runs)
                    for blip in para._element.xpath('.//a:blip'):
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.rels:
                            rel = doc.part.rels[rId]
                            if "image" in rel.target_ref:
                                img_ext = rel.target_part.content_type.split('/')[-1]
                                img_name = f"img_{uuid.uuid4().hex[:6]}.{img_ext}"
                                with open(os.path.join(IMAGE_DIR, img_name), "wb") as f: f.write(rel.target_part.blob)
                                para_text += f"\n[图片:{img_name}]\n"
                    if para_text.strip(): text_content.append(para_text)
            return "\n".join(text_content)
        except: return None
    return None

def set_run_font(run, font_name, size_pt, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color: run.font.color.rgb = color

def sanitize_image_for_docx(img_path):
    try:
        clean_path = img_path + "_v3_super_clean.png" 
        
        if not os.path.exists(clean_path):
            with Image.open(img_path) as img:
                if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                    bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
                    bg.paste(img, mask=img.convert('RGBA').split()[3])
                    img = bg.convert("RGB")
                else:
                    img = img.convert("RGB")
                    
                gray = img.convert("L")
                img_arr = np.array(gray)
                
                if np.mean(img_arr) < 90:
                    img_arr = 255 - img_arr
                    
                blur_img = Image.fromarray(img_arr).filter(ImageFilter.GaussianBlur(radius=30))
                blur_arr = np.array(blur_img, dtype=np.float32) + 1e-5 
                img_arr_f = np.array(img_arr, dtype=np.float32)
                
                norm_arr = np.clip((img_arr_f / blur_arr) * 255.0, 0, 255)
                norm_arr = np.where(norm_arr > 200, 255, norm_arr) 
                norm_arr = np.where(norm_arr < 150, norm_arr * 0.7, norm_arr) 
                norm_arr = np.clip(norm_arr, 0, 255).astype(np.uint8)
                
                final_img = Image.fromarray(norm_arr)
                
                inv_for_bbox = ImageOps.invert(final_img)
                bbox = inv_for_bbox.getbbox()
                if bbox:
                    margin = 10
                    left = max(0, bbox[0] - margin)
                    top = max(0, bbox[1] - margin)
                    right = min(final_img.width, bbox[2] + margin)
                    bottom = min(final_img.height, bbox[3] + margin)
                    final_img = final_img.crop((left, top, right, bottom))
                    
                final_img.save(clean_path, "PNG")
        return clean_path
    except Exception as e:
        print(f"Sanitize error: {e}")
        return img_path 

def generate_word_direct(sections, show_answer, paper_size, meta_info):
    doc = docx.Document()
    section_cfg = doc.sections[0]
    
    if paper_size == "A3 双栏正式版":
        section_cfg.page_width = Mm(420)
        section_cfg.page_height = Mm(297)
        sectPr = section_cfg._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else: cols = cols[0]
        cols.set(qn('w:num'), '2')     
        cols.set(qn('w:space'), '720') 
        max_img_width = Cm(7.5)
    else:
        section_cfg.page_width = Mm(210)
        section_cfg.page_height = Mm(297)
        max_img_width = Cm(14)

    header = section_cfg.header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    has_logo_or_name = False
    
    if meta_info.get("logo_bytes"):
        run_logo = p_header.add_run()
        try:
            run_logo.add_picture(io.BytesIO(meta_info["logo_bytes"]), height=Cm(1.5))
            has_logo_or_name = True
        except: pass
            
    if meta_info.get("school_name"):
        run_name = p_header.add_run(f" {meta_info['school_name']}")
        set_run_font(run_name, '黑体', 15, bold=True)
        has_logo_or_name = True
        
    if meta_info.get("watermark"):
        if has_logo_or_name:
            p_header.add_run("\n")
        run_wm = p_header.add_run("机密 ★ 内部考卷文件")
        set_run_font(run_wm, '黑体', 14, bold=True, color=RGBColor(211, 211, 211))

    footer = section_cfg.footer
    p_f = footer.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p_f.add_run("第 "); set_run_font(run1, '宋体', 10)
    run2 = p_f.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr1 = OxmlElement('w:instrText'); instr1.set(qn('xml:space'), 'preserve'); instr1.text = 'PAGE \\* MERGEFORMAT'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run2._r.extend([fld1, instr1, fld2, fld3])
    run3 = p_f.add_run(" 页  共 "); set_run_font(run3, '宋体', 10)
    run4 = p_f.add_run()
    fld4 = OxmlElement('w:fldChar'); fld4.set(qn('w:fldCharType'), 'begin')
    instr2 = OxmlElement('w:instrText'); instr2.set(qn('xml:space'), 'preserve'); instr2.text = 'NUMPAGES \\* MERGEFORMAT'
    fld5 = OxmlElement('w:fldChar'); fld5.set(qn('w:fldCharType'), 'separate')
    fld6 = OxmlElement('w:fldChar'); fld6.set(qn('w:fldCharType'), 'end')
    run4._r.extend([fld4, instr2, fld5, fld6])
    run5 = p_f.add_run(" 页"); set_run_font(run5, '宋体', 10)

    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(0)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    set_run_font(p_title.add_run('《建筑材料》智能化测试卷'), '黑体', 18, bold=True)
    
    if meta_info.get('subtitle'):
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(6)
        set_run_font(p_sub.add_run(meta_info['subtitle']), '黑体', 14)
        
    if meta_info.get('instructions'):
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_inst.paragraph_format.left_indent = Cm(0.6)
        p_inst.paragraph_format.space_after = Pt(6)
        set_run_font(p_inst.add_run("【考生须知】：\n" + meta_info['instructions']), '楷体', 10.5)
    
    meta_text = f"（考试时间：{meta_info['time']}分钟    满分：{meta_info['score']:g}分    考试形式：{meta_info['type']}）"
    p_meta = doc.add_paragraph(meta_text)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(12)
    set_run_font(p_meta.runs[0], '宋体', 10.5)

    p_info = doc.add_paragraph('班级：__________  姓名：__________  学号：__________  得分：__________')
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(18)
    set_run_font(p_info.runs[0], '宋体', 12)
    
    def render_question(q, doc, index, parent_id=""):
        q_type = normalize_type(q.get('type', ''))
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8) 
        p_q.paragraph_format.left_indent = Cm(0.6) if not parent_id else Cm(1.2)
        p_q.paragraph_format.first_line_indent = Cm(-0.6)
        
        raw_content = str(q.get('content', ''))
        img_tags = re.findall(r'\[图片:(.*?)\]', raw_content)
        clean_text = re.sub(r'\[图片:.*?\]', '', raw_content).strip()
        clean_text = clean_stem(clean_text)
        
        ans_raw = str(q.get('answer', '')).strip()
        if not ans_raw or ans_raw.lower() in ['null', 'none', '']: ans_raw = "待补全"
            
        q_id_val = str(q.get('id', '')).strip()
        q_id = q_id_val if q_id_val else str(index+1)
        clean_text = re.sub(r'^\d+[\.\、\s]+', '', clean_text)
        clean_text = clean_text.replace("$$", "").replace("$", "")
        
        if q_type in ['单选题', '多选题', '判断题']:
            ans_text = ""
            if show_answer:
                if q_type == '判断题' and ans_raw not in ["待补全"]: ans_text = '√' if ans_raw in ['对', '正确', '√', 'T', '1'] else '×'
                else: ans_text = ans_raw
                
            matches = list(re.finditer(r'[\(（]\s*[\)）]', clean_text))
            if matches:
                last_match = matches[-1]
                start, end = last_match.span()
                run_prefix = p_q.add_run(f"{q_id}. {clean_text[:start]}")
                set_run_font(run_prefix, '宋体', 12)
                set_run_font(p_q.add_run("（"), '宋体', 12)
                if show_answer and ans_text:
                    # ================= V2.9 扩宽括号内部的答案预留空间 =================
                    run_ans = p_q.add_run(f"  {ans_text}  "); set_run_font(run_ans, '宋体', 12, bold=True, color=RGBColor(255, 0, 0))
                else: 
                    p_q.add_run("      ") # 从 3 个空格加长到了 6 个空格，符合排版国标
                set_run_font(p_q.add_run("）"), '宋体', 12)
                if clean_text[end:]: set_run_font(p_q.add_run(clean_text[end:]), '宋体', 12)
            else:
                set_run_font(p_q.add_run(f"{q_id}. {clean_text}"), '宋体', 12)
                set_run_font(p_q.add_run(" （"), '宋体', 12)
                if show_answer and ans_text:
                    run_ans = p_q.add_run(f"  {ans_text}  "); set_run_font(run_ans, '宋体', 12, bold=True, color=RGBColor(255, 0, 0))
                else: 
                    p_q.add_run("      ") # 从 3 个空格加长到了 6 个空格，符合排版国标
                set_run_font(p_q.add_run("）"), '宋体', 12)
        else:
            set_run_font(p_q.add_run(f"{q_id}. {clean_text}"), '宋体', 12)
            if q_type == '填空题' and show_answer:
                run_ans = p_q.add_run(f"  [答案：{ans_raw}]"); set_run_font(run_ans, '宋体', 12, color=RGBColor(255, 0, 0))
            
        for img_name in img_tags:
            img_path = os.path.join(IMAGE_DIR, img_name)
            p_img = doc.add_paragraph()
            p_img.paragraph_format.left_indent = Cm(0.6) if not parent_id else Cm(1.2)
            run_img = p_img.add_run()
            
            if os.path.exists(img_path):
                try:
                    safe_img_path = sanitize_image_for_docx(img_path)
                    with Image.open(safe_img_path) as tmp_img:
                        w_px, h_px = tmp_img.size
                        ratio = w_px / float(h_px) if h_px > 0 else 1.0
                        
                    if ratio <= 1.3:
                        target_w = min(w_px * 0.02, 4.0)
                    elif ratio >= 2.0:
                        target_w = min(w_px * 0.02, max_img_width.cm)
                    else:
                        target_w = min(w_px * 0.02, 7.0)
                        
                    target_w = min(target_w, max_img_width.cm)
                    target_w = max(target_w, 1.5)
                    
                    run_img.add_picture(safe_img_path, width=Cm(target_w))
                except Exception as e:
                    run_img.text = f"\n[❌ 引擎将图片写入Word失败，原图可能损坏: {img_name}]"
                    set_run_font(run_img, '黑体', 11, color=RGBColor(255, 0, 0), bold=True)
            else:
                run_img.text = f"\n[⚠️ 严重警告：服务器本地找不到图片 {img_name}！请确保在上传原卷后再生成！]"
                set_run_font(run_img, '黑体', 11, color=RGBColor(255, 0, 0), bold=True)
        
        layout = q.get('layout', 'none')
        raw_opts = q.get('options', [])
        if layout != 'none' and raw_opts:
            opts = [clean_option(opt) for opt in raw_opts]
            tab_spacing = 2.5 if paper_size == "A3 双栏正式版" else 3.0 
            indent_cm = 0.6 if not parent_id else 1.2
            if layout == 'inline' and len(opts) == 4:
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.space_before = Pt(2); p_opt.paragraph_format.left_indent = Cm(indent_cm)
                p_opt.paragraph_format.tab_stops.add_tab_stop(Cm(indent_cm + tab_spacing))
                p_opt.paragraph_format.tab_stops.add_tab_stop(Cm(indent_cm + tab_spacing * 2))
                p_opt.paragraph_format.tab_stops.add_tab_stop(Cm(indent_cm + tab_spacing * 3))
                set_run_font(p_opt.add_run(f"A. {opts[0]}\tB. {opts[1]}\tC. {opts[2]}\tD. {opts[3]}"), '宋体', 12)
            elif layout == 'grid' and len(opts) == 4:
                p_opt1 = doc.add_paragraph(); p_opt1.paragraph_format.space_before = Pt(2); p_opt1.paragraph_format.left_indent = Cm(indent_cm)
                p_opt1.paragraph_format.tab_stops.add_tab_stop(Cm(indent_cm + tab_spacing * 2))
                set_run_font(p_opt1.add_run(f"A. {opts[0]}\tB. {opts[1]}"), '宋体', 12)
                p_opt2 = doc.add_paragraph(); p_opt2.paragraph_format.left_indent = Cm(indent_cm)
                p_opt2.paragraph_format.tab_stops.add_tab_stop(Cm(indent_cm + tab_spacing * 2))
                set_run_font(p_opt2.add_run(f"C. {opts[2]}\tD. {opts[3]}"), '宋体', 12)
            else: 
                for j, opt in enumerate(opts):
                    p_opt_block = doc.add_paragraph(f"{chr(65+j)}. {opt}")
                    p_opt_block.paragraph_format.left_indent = Cm(indent_cm)
                    if j == 0: p_opt_block.paragraph_format.space_before = Pt(2)
                    set_run_font(p_opt_block.runs[0], '宋体', 12)
        
        if q_type == '绘图题':
            p_draw = doc.add_paragraph()
            p_draw.paragraph_format.left_indent = Cm(0.6)
            run_draw = p_draw.add_run("（请在此方框内作图）\n")
            set_run_font(run_draw, '宋体', 10, color=RGBColor(128,128,128))
            for _ in range(8): doc.add_paragraph()

        if q_type in ['简答题', '计算题', '论述题', '分析题'] and q_type != '绘图题':
            if show_answer:
                p_ans = doc.add_paragraph(); p_ans.paragraph_format.space_before = Pt(6)
                set_run_font(p_ans.add_run(f"【参考答案】：{ans_raw}"), '宋体', 12, color=RGBColor(255, 0, 0))
                doc.add_paragraph()
            else:
                for _ in range(8 if q_type == '计算题' else 5): doc.add_paragraph()
                    
        explanation = str(q.get('explanation', '')).strip()
        if show_answer and explanation and explanation.lower() not in ['null', 'none']:
            p_exp = doc.add_paragraph(); p_exp.paragraph_format.left_indent = Cm(0.6); p_exp.paragraph_format.space_before = Pt(2); p_exp.paragraph_format.space_after = Pt(6)
            set_run_font(p_exp.add_run("【解析】："), '黑体', 10.5, color=RGBColor(255, 0, 0))
            clean_exp = re.sub(r'<[^>]+>', '', explanation).strip().replace('解析：', '').replace('答案：', '').strip()
            set_run_font(p_exp.add_run(clean_exp), '宋体', 10.5, color=RGBColor(255, 0, 0))
            
        if 'sub_questions' in q and isinstance(q['sub_questions'], list):
            for sub_i, sub_q in enumerate(q['sub_questions']):
                render_question(sub_q, doc, sub_i, parent_id=q_id)

    for section in sections:
        p_sec = doc.add_paragraph(); p_sec.paragraph_format.space_before = Pt(12); p_sec.paragraph_format.space_after = Pt(6)
        set_run_font(p_sec.add_run(section['title']), '黑体', 14, bold=True)
        for i, q in enumerate(section['questions']):
            render_question(q, doc, i)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def robust_json_parse(json_text):
    clean_json = re.sub(r'^```(json)?\s*', '', json_text.strip(), flags=re.IGNORECASE)
    clean_json = re.sub(r'\s*```$', '', clean_json, flags=re.IGNORECASE).strip()
    if clean_json.startswith('{') and clean_json.endswith(']'): clean_json = '[' + clean_json
    elif clean_json.startswith('[') and clean_json.endswith('}'): clean_json = clean_json + ']'
    elif clean_json.startswith('{') and clean_json.endswith('}'): clean_json = '[' + clean_json + ']'
    
    qs = json.loads(clean_json, strict=False)
    if isinstance(qs, dict): qs = [qs]
    for q in qs:
        ans = q.get('answer', '')
        if isinstance(ans, list): q['answer'] = "".join(str(x) for x in ans)
        elif ans is None: q['answer'] = ""
        else: q['answer'] = str(ans).strip()
    return qs

tab1, tab2, tab3 = st.tabs(["⚡ 全自动智能解析", "🌉 桥接人机协作", "🛒 试卷排版组装 & 数据导出"])

with tab1:
    col_upload, col_paste = st.columns(2)
    with col_upload:
        st.markdown("##### 方法 1：上传文件解析 (可多卷混编)")
        uploaded_auto = st.file_uploader("支持 PDF/Word/图片/Txt", type=["png", "jpg", "jpeg", "docx", "pdf", "doc", "txt"], key="auto_up", label_visibility="collapsed")
    with col_paste:
        st.markdown("##### 方法 2：直接粘贴纯文本")
        pasted_text = st.text_area("在此粘贴 Markdown 或纯文本内容", height=100, label_visibility="collapsed")
    
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🚀 一键提取并进入手术台", type="primary", use_container_width=True):
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key or "your_api_key" in api_key:
            st.error("未检测到有效 API Key，请去 .env 配置。")
        elif not uploaded_auto and not pasted_text.strip():
            st.warning("⚠️ 请先上传文件，或者粘贴试卷内容！")
        else:
            with st.spinner("正在提取内容，并呼叫 Gemini 进行多模态解析..."):
                raw_text = ""
                if uploaded_auto: raw_text = extract_text_and_images(uploaded_auto)
                elif pasted_text.strip(): raw_text = pasted_text.strip()
                    
                if raw_text:
                    try:
                        api_base = os.getenv("GEMINI_API_BASE")
                        if api_base:
                            genai.configure(api_key=api_key, client_options={"api_endpoint": api_base})
                        else:
                            genai.configure(api_key=api_key)
                        
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        prompt = f"""你是一个高级试卷排版员。将以下文本整理为严格的JSON数组。
包含字段: id(必须提取原卷题号), type(单选/多选/判断/填空/简答/计算/绘图), section(所属大题板块), content(题干), options(字符串数组), answer(答案), explanation(详细解析)。
【极其重要1】：如果文本中存在 `[图片:img_xxxx.png]`，必须原封不动保留在 content 中！
【极其重要2】：如果是阅读理解或共用题干的“套题”，请在父题中只保留背景材料作为 content，并将下面的小题存放在父题的 `sub_questions` 字段中。
文本：{raw_text}"""
                        response = model.generate_content(prompt)
                        questions = robust_json_parse(response.text)
                        for q in questions:
                            q['layout'] = calculate_option_layout(q)
                            st.session_state.cart.append(q)
                        st.success(f"🎉 成功解析 {len(questions)} 道大题！请前往第三页核解。")
                    except Exception as e:
                        st.error(f"解析失败：{e}")

with tab2:
    col_step1, col_step2 = st.columns(2)
    with col_step1:
        st.markdown("##### 步骤 1：生成大模型 Prompt")
        uploaded_bridge = st.file_uploader("上传试卷", type=["png", "jpg", "jpeg", "docx", "pdf", "doc", "txt"], key="b_up")
        pasted_bridge_text = st.text_area("或者在此直接粘贴文本：", height=100, key="bridge_paste", label_visibility="collapsed")
        
        if st.button("生成提示词"):
            raw_text = ""
            if uploaded_bridge: raw_text = extract_text_and_images(uploaded_bridge)
            elif pasted_bridge_text.strip(): raw_text = pasted_bridge_text.strip()
                
            if raw_text:
                st.success("复制下方提示词发送给大模型：")
                st.code(f"你是一个高级排版员。请将以下文字整理为严格JSON数组。字段：id, type, section, content, options, answer, explanation, sub_questions(用于阅读套题)。\n{raw_text}", language="text")
            else: st.warning("请先上传文件或粘贴内容！")
                
    with col_step2:
        st.markdown("##### 步骤 2：导入 JSON 代码")
        json_input = st.text_area("粘贴JSON：", height=150, label_visibility="collapsed")
        if st.button("校验并加入试题库", type="primary"):
            try:
                qs = robust_json_parse(json_input)
                for q in qs:
                    q['layout'] = calculate_option_layout(q)
                    st.session_state.cart.append(q)
                st.success(f"✅ 成功加入 {len(qs)} 道题目。")
            except Exception as e:
                st.error(f"JSON 解析依然失败，可能存在括号未闭合等严重语法错误：{e}")

with tab3:
    current_sections = []
    for q in st.session_state.cart:
        sec_raw = str(q.get('section', '')).strip()
        type_raw = normalize_type(q.get('type', ''))
        
        if type_raw.replace('题', '') in sec_raw:
            sec_name = sec_raw
        elif sec_raw and sec_raw.lower() not in ['null', 'none']:
            sec_name = f"{sec_raw} - {type_raw}"
        else:
            sec_name = type_raw
            
        q['clean_section'] = sec_name
        if sec_name not in current_sections: current_sections.append(sec_name)
        
    st.session_state.section_order = [s for s in (st.session_state.section_order + [x for x in current_sections if x not in st.session_state.section_order]) if s in current_sections]

    section_configs = {}
    total_calc_score = 0
    total_chars = sum(len(str(q)) for q in st.session_state.cart)
    
    if len(st.session_state.cart) > 0:
        st.metric(label="当前手术台题量 (支持多卷混编购物车)", value=f"{len(st.session_state.cart)} 道")
        with st.expander("🛠️ 试题手术台", expanded=True):
            edited_cart = st.data_editor(st.session_state.cart, num_rows="dynamic", use_container_width=True, hide_index=True)
            col_save1, col_save2 = st.columns(2)
            with col_save1:
                if st.button("💾 保存手术台修改"):
                    st.session_state.cart = edited_cart
                    st.success("已生效！")
                    st.rerun()
            with col_save2:
                save_tag = st.text_input("给这批题打个标签再存入库 (如: 摸底考)", label_visibility="collapsed", placeholder="输入题库分类标签")
                if st.button("📥 存入本地知识库 (SQLite)"):
                    if not save_tag: st.warning("请填写标签")
                    else:
                        cnt = save_to_db(edited_cart, save_tag)
                        st.success(f"成功将 {cnt} 道题存入本地数据库！")

    with st.expander("📝 试卷表头、附加信息与排版设定"):
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            meta_time = st.number_input("考试时间 (分钟)", value=90, step=10)
            meta_subtitle = st.text_input("副标题 (可选)", placeholder="如：2026-2027学年上学期期中测验")
        with col_m2:
            meta_type = st.selectbox("考试形式", ["闭卷", "开卷", "半开卷"])
            meta_instructions = st.text_area("考生须知 (可选)", placeholder="如：请用黑色碳素笔作答...", height=68)
            
        st.markdown("---")
        col_logo1, col_logo2 = st.columns(2)
        with col_logo1:
            meta_logo_file = st.file_uploader("上传学校 Logo (PNG/JPG，居中)", type=["png", "jpg", "jpeg"], key="meta_logo")
        with col_logo2:
            meta_school_name = st.text_input("自定义学校名字 (显示在Logo后)", placeholder="如：四川省营山职业高级中学")
            st.markdown("<br>", unsafe_allow_html=True)
            meta_watermark = st.checkbox("生成防伪/保密页眉水印 (机密★内部文件)", value=False)
            
        if st.session_state.section_order:
            st.markdown("---")
            st.markdown("##### 📐 大题顺序与智能分值")
            for i, sec in enumerate(st.session_state.section_order):
                qs_in_sec = [q for q in st.session_state.cart if q['clean_section'] == sec]
                q_count = len(qs_in_sec)
                
                types_in_sec = "".join([str(q.get('type', '')) for q in qs_in_sec])
                combined_hint = (sec + types_in_sec).lower()
                
                if '多选' in combined_hint or 'multiple' in combined_hint: default_pt = 4.0
                elif '单选' in combined_hint or 'single' in combined_hint: default_pt = 3.0
                elif '判断' in combined_hint or 'judge' in combined_hint or 'true' in combined_hint or 'false' in combined_hint: default_pt = 2.0
                elif '填空' in combined_hint or 'fill' in combined_hint: default_pt = 2.0
                else: default_pt = 5.0
                
                col_n, col_up, col_down, col_s = st.columns([4, 0.5, 0.5, 2])
                with col_n: st.markdown(f"**{sec}** *(共 {q_count} 题)*")
                with col_up:
                    if st.button("⬆️", key=f"up_{sec}", disabled=(i==0)):
                        st.session_state.section_order[i], st.session_state.section_order[i-1] = st.session_state.section_order[i-1], st.session_state.section_order[i]; st.rerun()
                with col_down:
                    if st.button("⬇️", key=f"down_{sec}", disabled=(i==len(st.session_state.section_order)-1)):
                        st.session_state.section_order[i], st.session_state.section_order[i+1] = st.session_state.section_order[i+1], st.session_state.section_order[i]; st.rerun()
                with col_s: score = st.number_input(f"单题分值", value=default_pt, step=0.5, min_value=0.5, key=f"scr_{sec}", label_visibility="collapsed")
                section_configs[sec] = {'score': score}
                total_calc_score += score * q_count
            st.info(f"💯 当前试卷满分为 **{total_calc_score:g} 分**")
    
    col1, col2 = st.columns(2)
    with col1:
        doc_type = st.radio("输出内容版本", ["纯净学生版", "教师解析版"])
    with col2:
        paper_size = st.radio("纸张规格", ["A4 单栏日常版", "A3 双栏正式版"])
        
    st.markdown("<br>", unsafe_allow_html=True)
    
    est_pages = (total_chars / 1500) if paper_size == "A4 单栏日常版" else (total_chars / 3000)
    if len(st.session_state.cart) > 0 and est_pages > 4:
        st.warning(f"⚠️ **排版预警**：系统预估本试卷字数与题量较多，{paper_size} 下可能超出 4 页（约 {est_pages:.1f} 页）。建议删减题目或在 Word 中缩小行距。")
    
    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 3])
    with col_btn1:
        if st.button("🗑️ 清空题库"): st.session_state.cart = []; st.session_state.section_order = []; st.rerun()
    with col_btn2:
        if len(st.session_state.cart) > 0:
            df = pd.DataFrame(st.session_state.cart)
            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False)
            st.download_button("📊 导出为 Excel", data=excel_buffer.getvalue(), file_name="试题数据_导出.xlsx", mime="application/vnd.ms-excel")
    with col_btn3:
        if st.button("🖨️ 生成 教务处标准 Word 试卷", type="primary", use_container_width=True):
            if len(st.session_state.cart) == 0: st.warning("题库为空！")
            else:
                with st.spinner("正在融合高级排版指令..."):
                    sections = group_questions(st.session_state.cart, section_configs, st.session_state.section_order)
                    meta_info = {
                        "time": meta_time, "score": total_calc_score, "type": meta_type, 
                        "subtitle": meta_subtitle, "instructions": meta_instructions,
                        "school_name": meta_school_name,
                        "watermark": meta_watermark,
                        "logo_bytes": meta_logo_file.getvalue() if meta_logo_file else None
                    }
                    file_stream = generate_word_direct(sections, doc_type == "教师解析版", paper_size, meta_info)
                    st.success("✅ 生成完毕！")
                    st.download_button(label=f"📥 下载试卷.docx", data=file_stream, file_name=f"试卷_{'教师版' if doc_type == '教师解析版' else '学生版'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
